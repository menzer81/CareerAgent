"""DOCX/PDF resume export service (Sprint 4)."""

from __future__ import annotations

import textwrap
from pathlib import Path

from app.config import get_settings
from app.schemas.resume import ResumePlan


class ResumeExportService:
    def __init__(self) -> None:
        self.settings = get_settings()

    @staticmethod
    def _markdown_lines(markdown: str) -> list[str]:
        lines = [line.rstrip() for line in markdown.splitlines()]
        return lines if lines else [""]

    def save_docx(self, job_posting_id: int, plan: ResumePlan) -> Path:
        try:
            from docx import Document
        except ModuleNotFoundError as exc:
            raise RuntimeError("DOCX export requires python-docx to be installed.") from exc

        export_dir = self.settings.reports_dir / "resumes"
        export_dir.mkdir(parents=True, exist_ok=True)
        output_path = export_dir / f"job_{job_posting_id}_resume.docx"

        doc = Document()
        for raw in self._markdown_lines(plan.markdown):
            if raw.startswith("### "):
                doc.add_heading(raw[4:].strip(), level=3)
            elif raw.startswith("## "):
                doc.add_heading(raw[3:].strip(), level=2)
            elif raw.startswith("# "):
                doc.add_heading(raw[2:].strip(), level=1)
            elif raw.startswith("- "):
                doc.add_paragraph(raw[2:].strip(), style="List Bullet")
            elif raw:
                doc.add_paragraph(raw)
            else:
                doc.add_paragraph("")
        doc.save(output_path)
        return output_path

    def save_pdf(self, job_posting_id: int, plan: ResumePlan) -> Path:
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
        except ModuleNotFoundError as exc:
            raise RuntimeError("PDF export requires reportlab to be installed.") from exc

        export_dir = self.settings.reports_dir / "resumes"
        export_dir.mkdir(parents=True, exist_ok=True)
        output_path = export_dir / f"job_{job_posting_id}_resume.pdf"

        c = canvas.Canvas(str(output_path), pagesize=letter)
        width, height = letter
        y = height - 72
        left_margin = 72
        line_height = 14
        max_chars = 95

        for raw in self._markdown_lines(plan.markdown):
            text = raw.strip()
            if not text:
                y -= line_height
            else:
                if text.startswith("#"):
                    text = text.lstrip("#").strip().upper()
                if text.startswith("- "):
                    text = f"• {text[2:].strip()}"
                for wrapped in textwrap.wrap(text, width=max_chars) or [""]:
                    if y <= 72:
                        c.showPage()
                        y = height - 72
                    c.drawString(left_margin, y, wrapped)
                    y -= line_height

            if y <= 72:
                c.showPage()
                y = height - 72

        c.save()
        return output_path
